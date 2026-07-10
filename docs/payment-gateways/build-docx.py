#!/usr/bin/env python3
"""Build payment-gateway-field.docx from documentation content."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(__file__).parent / "payment-gateway-field.docx"

doc = Document()

doc.add_heading("Payment Gateway Field", 0)

meta = doc.add_paragraph()
meta.add_run("Category: Payment Gateways\n").italic = True
meta.add_run("Plugin: Everest Forms Pro\n").italic = True
meta.add_run("Since: 1.9.15\n").italic = True
meta.add_run("Publish URL: https://docs.everestforms.net/docs/payment-gateway-field/\n").italic = True

doc.add_paragraph(
    "The Payment Gateway field lets visitors choose how they pay on a single form—Stripe, PayPal, Square, "
    "Authorize.Net, Mollie, or Razorpay—without building separate forms or wiring Conditional Logic between "
    "payment methods."
)
doc.add_paragraph(
    "This replaces the older workflow where you added a Multiple Choice field and used Conditional Logic on the "
    "Payments tab and Credit Card field. With the Payment Gateway field, one drag-and-drop field handles gateway "
    "selection, card UI, and redirect messaging."
)

doc.add_heading("Prerequisites", level=2)
for item in [
    "Everest Forms Pro",
    "At least one payment gateway add-on installed and activated",
    "Global credentials configured for each gateway you want to offer",
    "SSL (HTTPS) recommended for card-based gateways (Stripe, Square, Authorize.Net)",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("What's New in Recent Versions?", level=2)
doc.add_paragraph(
    "With Everest Forms Pro 1.9.15+, the Payment Gateway field is available under Payment Fields in the form builder."
)
t = doc.add_table(rows=5, cols=2)
t.style = "Table Grid"
t.rows[0].cells[0].text = "Before (legacy)"
t.rows[0].cells[1].text = "With Payment Gateway field"
for i, (a, b) in enumerate(
    [
        ("Multiple Choice for PayPal / Stripe", "Dedicated Payment Gateway field with branded cards"),
        (
            "Enable each gateway on the Payments tab",
            "Gateways appear when the add-on is active and credentials are saved",
        ),
        (
            "Conditional Logic on Payments tab + Credit Card field",
            "Card fields show/hide automatically when the user picks a gateway",
        ),
        (
            "One Credit Card field per Stripe form",
            "Stripe, Square, and Authorize.Net card UI mounts inside the selector when needed",
        ),
    ],
    1,
):
    t.rows[i].cells[0].text = a
    t.rows[i].cells[1].text = b
doc.add_paragraph(
    "One field per form: Only one Payment Gateway field can be added to a form. After you add it, legacy Credit "
    "Card, Square Payment, and Authorize.Net fields are hidden from the field list to avoid duplicate payment UIs."
)

doc.add_heading("Supported Payment Gateways", level=2)
t2 = doc.add_table(rows=7, cols=3)
t2.style = "Table Grid"
t2.rows[0].cells[0].text = "Gateway"
t2.rows[0].cells[1].text = "Add-on"
t2.rows[0].cells[2].text = "Checkout type"
for i, row in enumerate(
    [
        ("Stripe", "Everest Forms Stripe", "On-page card (Stripe Elements)"),
        ("PayPal Standard", "Everest Forms PayPal Standard", "Redirect to PayPal"),
        ("Square", "Everest Forms Square (Pro)", "On-page card (Square Web Payments SDK)"),
        ("Authorize.Net", "Everest Forms Authorize.Net", "On-page card (Accept.js)"),
        ("Mollie", "Everest Forms Mollie (Pro)", "Redirect to Mollie"),
        ("Razorpay", "Everest Forms Razorpay", "Redirect / Razorpay checkout"),
    ],
    1,
):
    for j, val in enumerate(row):
        t2.rows[i].cells[j].text = val
doc.add_paragraph("A gateway is selectable on the frontend only when:")
for item in [
    "Its add-on is active",
    "Global credentials are configured (PayPal can also use per-form email when global settings are disabled)",
    "The gateway is enabled in the Payment Gateway field Gateways list",
]:
    doc.add_paragraph(item, style="List Number")

doc.add_heading("Installation", level=2)
for i, step in enumerate(
    [
        "Purchase Everest Forms Pro from WPEverest.",
        "Install and activate Everest Forms Pro from your WordPress dashboard.",
        "Install and activate payment add-ons from Everest Forms → Add-ons.",
        "Configure each gateway under Everest Forms → Settings → Payments.",
    ],
    1,
):
    doc.add_paragraph(f"{i}. {step}", style="List Number")
note = doc.add_paragraph()
note.add_run(
    "For a detailed guide, read our documentation on how to install and activate Everest Forms Pro."
).italic = True

doc.add_heading("Configure Global Payment Credentials", level=2)
doc.add_paragraph(
    "The Payment Gateway field does not replace global payment settings. Each add-on still needs API keys, "
    "merchant email, or tokens saved once site-wide."
)
doc.add_paragraph("Navigate to Everest Forms → Settings → Payments and configure:")
t3 = doc.add_table(rows=7, cols=2)
t3.style = "Table Grid"
t3.rows[0].cells[0].text = "Gateway"
t3.rows[0].cells[1].text = "Settings to complete"
for i, (a, b) in enumerate(
    [
        ("Stripe", "Publishable key and Secret key (Live/Test)"),
        ("PayPal Standard", "PayPal email, Mode (Sandbox/Production), Payment type"),
        ("Square", "Application ID, Access token, Location ID"),
        ("Authorize.Net", "API Login ID and Transaction Key"),
        ("Mollie", "API key"),
        ("Razorpay", "Key ID and Secret"),
    ],
    1,
):
    t3.rows[i].cells[0].text = a
    t3.rows[i].cells[1].text = b
doc.add_paragraph("Also set Currency under the General payments section.")
doc.add_paragraph("Gateway guides: docs.everestforms.net/docs/paypal-standard/, /stripe/, /square/, /authorize-net/, /mollie/, /razorpay/")

doc.add_heading("Add the Payment Gateway Field to Your Form", level=2)
for i, step in enumerate(
    [
        "Open Everest Forms → All Forms and edit your form (or create a new one).",
        "In the builder, open Fields → Payment Fields.",
        "Drag Payment Gateway into the form.",
        "Add payment line items (Single Item, Multiple Choice, Total, etc.).",
        "Click Save.",
    ],
    1,
):
    doc.add_paragraph(f"{i}. {step}", style="List Number")
doc.add_paragraph("The field is required by default so visitors must pick a payment method before submitting.")

doc.add_heading("Choose Which Gateways Appear on the Form", level=2)
for i, step in enumerate(
    [
        "Click the Payment Gateway field in the builder.",
        "In Field Options on the left, open Gateways.",
        "Use the toggle to include or exclude each gateway; drag to reorder.",
        "Expand a gateway row (chevron) for per-gateway settings.",
    ],
    1,
):
    doc.add_paragraph(f"{i}. {step}", style="List Number")
doc.add_paragraph("Notes:", style="List Bullet")
for note in [
    "Gateways without credentials show as disabled until connected globally.",
    "PayPal is unchecked by default if no global PayPal email is saved.",
    "If every toggle is off, the form prompts you to enable at least one gateway.",
]:
    doc.add_paragraph(note, style="List Bullet")

doc.add_heading("Per-Gateway Settings in the Form Builder", level=2)
t4 = doc.add_table(rows=6, cols=2)
t4.style = "Table Grid"
t4.rows[0].cells[0].text = "Gateway"
t4.rows[0].cells[1].text = "Typical accordion options"
for i, (a, b) in enumerate(
    [
        ("PayPal", "Use global settings, email, Sandbox/Production, Cancel URL"),
        ("Stripe", "Recurring toggle, iDEAL, map customer fields"),
        ("Authorize.Net", "Map Customer Email (required for subscriptions)"),
        ("Square", "Subscription options with Subscription Plan field"),
        ("Mollie", "Subscription description and recurring defaults"),
    ],
    1,
):
    t4.rows[i].cells[0].text = a
    t4.rows[i].cells[1].text = b
doc.add_paragraph(
    "You do not need Enable PayPal / Enable Stripe on the legacy Payments tab when using the Payment Gateway field."
)

doc.add_heading("More on Setup and Configuration", level=2)

doc.add_heading("Payment Fields", level=3)
doc.add_paragraph("Single Item item types:")
for item in [
    "Pre-Defined: fixed price on the frontend",
    "User Defined: visitor enters amount (donations)",
    "Hidden: price not shown; configured price is charged",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph(
    "Use Multiple Choice, Checkbox, Quantity, and Total to build multi-item carts."
)

doc.add_heading("Build a One-Time Payment Form", level=3)
for i, step in enumerate(
    [
        "Add Payment Gateway; enable Stripe, PayPal, or others.",
        "Add Single Item with your price.",
        "Optionally add Total and contact fields.",
        "Save and embed via shortcode or block.",
    ],
    1,
):
    doc.add_paragraph(f"{i}. {step}", style="List Number")
doc.add_paragraph("Stripe / Square / Authorize.Net: card fields appear on selection.")
doc.add_paragraph("PayPal / Mollie / Razorpay: redirect message; user pays on provider site.")

doc.add_heading("Subscription Plan on Form", level=3)
for i, step in enumerate(
    [
        "Add Payment Gateway; enable subscription-capable gateways.",
        "Add Subscription Plan; configure price, period, trial, expiry.",
        "Map gateway-specific options (e.g. Authorize.Net Customer Email).",
        "Save the form.",
    ],
    1,
):
    doc.add_paragraph(f"{i}. {step}", style="List Number")
doc.add_paragraph(
    "With Payment Gateway + Subscription Plan, subscription mode runs without the legacy Payments tab recurring toggle."
)

doc.add_heading("Frontend View", level=2)
for item in [
    "Card grid of payment logos on the live form.",
    "One gateway required unless only one is enabled (auto-selected).",
    "Card gateways reveal inline card UI; redirect gateways show a secure redirect message.",
    "Works with shortcode, Gutenberg block, and AJAX submission.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Legacy Payments Tab vs Payment Gateway Field", level=2)
t5 = doc.add_table(rows=5, cols=2)
t5.style = "Table Grid"
t5.rows[0].cells[0].text = "Scenario"
t5.rows[0].cells[1].text = "Recommended approach"
for i, (a, b) in enumerate(
    [
        ("One gateway only, simple form", "Legacy Payments tab"),
        ("Visitor chooses PayPal or Stripe (or more)", "Payment Gateway field"),
        ("Subscriptions with plan choices", "Subscription Plan + Payment Gateway"),
        ("Old forms with Conditional Logic", "Keep working; migrate when convenient"),
    ],
    1,
):
    t5.rows[i].cells[0].text = a
    t5.rows[i].cells[1].text = b
doc.add_paragraph(
    "Do not mix Payment Gateway with standalone Credit Card, Square Payment, or Authorize.Net fields."
)

doc.add_heading("How to add Stripe as a payment option with PayPal? (legacy)", level=3)
doc.add_paragraph(
    "Older forms used Multiple Choice + Conditional Logic. New forms should use the Payment Gateway field. "
    "See PayPal Standard and Stripe documentation on docs.everestforms.net for legacy steps."
)

doc.add_heading("Troubleshooting", level=2)
issues = [
    (
        "No payment methods appear on the form",
        [
            "Activate the payment add-on under Everest Forms → Add-ons.",
            "Save global credentials under Settings → Payments.",
            "Enable at least one gateway toggle in field options.",
            "For PayPal, set global email or per-form email in accordion.",
        ],
    ),
    (
        "Please choose a valid payment method on submit",
        [
            "Selected gateway must be in allowlist and connected.",
            "User must select a valid gateway radio option.",
        ],
    ),
    (
        "PayPal shows Things don't appear to be working",
        [
            "Match Sandbox mode to sandbox business email.",
            "Ensure form total is greater than zero.",
            "Check Everest Forms → Tools → Payment Log.",
        ],
    ),
    (
        "Stripe, Square, or Authorize.Net card form does not show",
        [
            "Use HTTPS on the form page.",
            "Confirm gateway toggle is on.",
            "Check browser console for JS/cache conflicts.",
        ],
    ),
    (
        "Authorize.Net subscriptions fail",
        [
            "Map Customer Email (Subscriptions) in Authorize.Net accordion to form Email field.",
        ],
    ),
]
for title, bullets in issues:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Finishing up", level=2)
doc.add_paragraph("Click Save in the form builder.")
doc.add_paragraph("Embed using the shortcode or Everest Forms block.")
doc.add_paragraph("Payment entries appear under Everest Forms → Entries.")

footer = doc.add_paragraph()
footer.add_run("Category: Payment Gateways · Last updated: June 2026").italic = True

doc.save(OUT)
print(f"Created: {OUT}")
